"""
Generate project improvement report as .docx
SaaS UMKM — Feature & Workflow Documentation
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

PROJECT_ROOT = r"C:\saasumkm"
SCREENSHOTS_DIR = os.path.join(PROJECT_ROOT, "tests", "screenshots")
OUTPUT_PATH = os.path.join(PROJECT_ROOT, "docs", "Laporan-SaaS-UMKM-v2.docx")

# Screenshot metadata — category, filename, caption
SECTIONS = [
    {
        "title": "1. Halaman Publik (Desktop)",
        "description": "Halaman yang dapat diakses tanpa login — landing page, autentikasi, dan storefront.",
        "shots": [
            ("desktop-public/01-landing.png", "Landing Page — halaman utama platform"),
            ("desktop-public/02-login.png", "Halaman Login"),
            ("desktop-public/03-register.png", "Halaman Registrasi Tenant Baru"),
            ("desktop-public/04-storefront.png", "Storefront Tenant — katalog produk publik"),
            ("desktop-public/05-storefront-cart.png", "Keranjang Belanja (Cart Sheet)"),
            ("desktop-public/06-track.png", "Tracking Pesanan (tanpa login)"),
        ],
    },
    {
        "title": "2. Dashboard (Desktop — Authenticated)",
        "description": "Panel manajemen bisnis untuk pemilik UMKM setelah login.",
        "shots": [
            ("desktop-auth/10-dashboard-home.png", "Dashboard Home — ringkasan bisnis"),
            ("desktop-auth/11-dashboard-insights.png", "Business Intelligence (RFM, CLV, Forecast)"),
            ("desktop-auth/12-dashboard-analytics.png", "Analytics — grafik penjualan"),
            ("desktop-auth/13-dashboard-orders.png", "Manajemen Pesanan"),
            ("desktop-auth/14-dashboard-order-detail.png", "Detail Pesanan + Status Transition"),
            ("desktop-auth/15-dashboard-products.png", "Katalog Produk"),
            ("desktop-auth/16-dashboard-product-new.png", "Form Tambah Produk Baru"),
            ("desktop-auth/17-dashboard-stock.png", "Manajemen Stok"),
            ("desktop-auth/18-dashboard-promos.png", "Promo & Voucher"),
            ("desktop-auth/19-dashboard-customers.png", "Database Pelanggan"),
            ("desktop-auth/20-dashboard-loyalty.png", "Loyalty Program (Poin & Tier)"),
            ("desktop-auth/21-dashboard-outlets.png", "Multi-Outlet Management"),
            ("desktop-auth/22-dashboard-staff.png", "Staff Management"),
            ("desktop-auth/23-dashboard-billing.png", "Billing & Subscription"),
            ("desktop-auth/24-dashboard-settings.png", "Pengaturan Toko"),
            ("desktop-auth/25-dashboard-pos.png", "Point of Sale (POS)"),
            ("desktop-auth/26-dashboard-theme-builder.png", "Theme Builder — kustomisasi branding"),
            ("desktop-auth/27-dashboard-custom-domain.png", "Custom Domain Setup"),
            ("desktop-auth/28-dashboard-template-builder.png", "Template Builder — layout storefront"),
            ("desktop-auth/29-dashboard-notifications.png", "Pusat Notifikasi"),
            ("desktop-auth/30-platform-admin.png", "Platform Admin Panel"),
            ("desktop-auth/31-dashboard-sidebar-collapsed.png", "Sidebar Collapsed (Icon Rail)"),
            ("desktop-auth/32-dashboard-ai-studio.png", "AI Studio — Content Generator"),
            ("desktop-auth/33-dashboard-logo-builder.png", "Logo Builder — parametric logo"),
        ],
    },
    {
        "title": "3. Mobile Responsive",
        "description": "Tampilan mobile (iPhone 14 Pro, 393×852) — memastikan UX optimal di perangkat genggam.",
        "shots": [
            ("mobile/01-landing.png", "Landing Page (Mobile)"),
            ("mobile/02-login.png", "Login (Mobile)"),
            ("mobile/03-storefront.png", "Storefront (Mobile)"),
            ("mobile/10-dashboard-home.png", "Dashboard Home (Mobile)"),
            ("mobile/11-dashboard-insights.png", "Insights (Mobile)"),
            ("mobile/12-dashboard-more-drawer.png", "Bottom Nav 'More' Drawer (Mobile)"),
        ],
    },
    {
        "title": "4. Cinematic Wide (1600px)",
        "description": "Tampilan layar lebar — memastikan layout tidak pecah di monitor besar.",
        "shots": [
            ("cinematic-wide/01-hero-cinematic.png", "Hero Section (Cinematic)"),
            ("cinematic-wide/02-landing-full.png", "Full Landing Page (Cinematic)"),
            ("cinematic-wide/03-pricing-band.png", "Pricing Band (Cinematic)"),
        ],
    },
]


def create_report():
    doc = Document()

    # -- Page margins --
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # -- Cover / Title --
    doc.add_paragraph()  # spacing
    title = doc.add_heading("Laporan Pengembangan", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading("SaaS UMKM — Platform Manajemen Bisnis", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"\nTanggal: {datetime.now().strftime('%d %B %Y')}\n").bold = False
    meta.add_run("Versi: 0.1.0\n")
    meta.add_run("Stack: Next.js 16 + React 19 + Prisma + TypeScript\n")

    doc.add_page_break()

    # -- Daftar Isi --
    doc.add_heading("Daftar Isi", level=1)
    toc_items = [
        "1. Halaman Publik (Desktop)",
        "2. Dashboard (Desktop — Authenticated)",
        "3. Mobile Responsive",
        "4. Cinematic Wide (1600px)",
        "5. Ringkasan Pengembangan",
        "6. Fitur Baru",
        "7. Bug Fix",
        "8. Test Coverage",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # -- Screenshot Sections --
    for section_data in SECTIONS:
        doc.add_heading(section_data["title"], level=1)
        doc.add_paragraph(section_data["description"])

        for filename, caption in section_data["shots"]:
            filepath = os.path.join(SCREENSHOTS_DIR, filename)
            if not os.path.exists(filepath):
                doc.add_paragraph(f"⚠️ Screenshot tidak ditemukan: {filename}")
                continue

            doc.add_paragraph()  # spacing
            doc.add_picture(filepath, width=Inches(6.0))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(f"Gambar: {caption}")
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_page_break()

    # -- Section 5: Ringkasan Pengembangan --
    doc.add_heading("5. Ringkasan Pengembangan", level=1)

    table = doc.add_table(rows=8, cols=3)
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Metrik", "Sebelum", "Sesudah"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    data = [
        ("Unit Tests", "20", "98 (+390%)"),
        ("Test Files", "5", "13"),
        ("Lib Coverage", "5/22 (23%)", "13/22 (59%)"),
        ("Bug Produksi Diperbaiki", "0", "3"),
        ("Fitur Baru", "0", "3"),
        ("CI Workflows", "4 ✅", "4 ✅"),
        ("Build Status", "✅", "✅"),
    ]
    for row_idx, (metric, before, after) in enumerate(data, start=1):
        table.rows[row_idx].cells[0].text = metric
        table.rows[row_idx].cells[1].text = before
        table.rows[row_idx].cells[2].text = after

    doc.add_paragraph()

    # -- Section 6: Fitur Baru --
    doc.add_heading("6. Fitur Baru", level=1)

    doc.add_heading("6.1 Loyalty Tiers (Bronze/Silver/Gold/Platinum)", level=2)
    doc.add_paragraph(
        "Sistem 4-tier berdasarkan total belanja kumulatif. "
        "Setiap tier memberi point multiplier saat earning — "
        "pelanggan loyal otomatis dapat poin lebih cepat."
    )
    tier_table = doc.add_table(rows=5, cols=4)
    tier_table.style = "Light Shading Accent 1"
    for i, h in enumerate(["Tier", "Min. Belanja", "Multiplier", "Perk"]):
        tier_table.rows[0].cells[i].text = h
    tier_data = [
        ("Bronze", "Rp 0", "1.0x", "Default"),
        ("Silver", "Rp 1.000.000", "1.25x", "+25% poin"),
        ("Gold", "Rp 5.000.000", "1.5x", "+50% poin + promo eksklusif"),
        ("Platinum", "Rp 20.000.000", "2.0x", "Poin 2x, prioritas layanan"),
    ]
    for r, (tier, spent, mult, perk) in enumerate(tier_data, start=1):
        tier_table.rows[r].cells[0].text = tier
        tier_table.rows[r].cells[1].text = spent
        tier_table.rows[r].cells[2].text = mult
        tier_table.rows[r].cells[3].text = perk

    doc.add_paragraph()
    doc.add_heading("6.2 Receipt Formatter (Struk WhatsApp)", level=2)
    doc.add_paragraph(
        "Endpoint /api/store/[slug]/receipt/[id] menghasilkan struk teks "
        "siap share via WhatsApp. Mendukung format monospace (```...```) "
        "untuk tampilan rapi di chat. Include: header tenant, items, "
        "subtotal/diskon/total, dan link storefront."
    )

    doc.add_heading("6.3 PPN Tax Calculator", level=2)
    doc.add_paragraph(
        "Kalkulator pajak PPN Indonesia sesuai UU HPP (11%). "
        "Mendukung mode EXCLUSIVE (pajak ditambah di atas) dan INCLUSIVE "
        "(pajak diekstrak dari total). Sudah ter-wire ke order response — "
        "tenant yang mengaktifkan taxEnabled mendapat breakdown PPN otomatis. "
        "Siap untuk tarif 12% di masa depan (configurable per-tenant)."
    )

    doc.add_page_break()

    # -- Section 7: Bug Fix --
    doc.add_heading("7. Bug Fix Kritis", level=1)

    doc.add_heading("7.1 Race Condition Overselling Stok", level=2)
    doc.add_paragraph("Severity: CRITICAL")
    doc.add_paragraph(
        "Masalah: Validasi stok dilakukan di luar transaksi database. "
        "Dua checkout bersamaan untuk unit terakhir bisa lolos validasi → "
        "stok menjadi negatif (oversold). Pola klasik check-then-act."
    )
    doc.add_paragraph(
        "Solusi: Atomic guarded decrement menggunakan updateMany dengan "
        "kondisi WHERE stock >= quantity. Jika stok sudah habis saat write, "
        "transaksi otomatis gagal dan order ditolak."
    )
    doc.add_paragraph("File: src/app/api/store/[slug]/order/route.ts")

    doc.add_heading("7.2 Loyalty Points Input Negatif", level=2)
    doc.add_paragraph("Severity: HIGH")
    doc.add_paragraph(
        "Masalah: calculatePoints(-50000) menghasilkan -5 poin. "
        "Input negatif atau NaN (refund, data corrupt) bisa mengurangi "
        "poin pelanggan secara diam-diam."
    )
    doc.add_paragraph(
        "Solusi: Guard Number.isFinite() dan <= 0 pada kedua fungsi "
        "calculatePoints() dan pointsToDiscount(). Input invalid → return 0."
    )
    doc.add_paragraph("File: src/lib/loyalty.ts")

    doc.add_heading("7.3 React title Array Error (SVG Tooltips)", level=2)
    doc.add_paragraph("Severity: MEDIUM")
    doc.add_paragraph(
        "Masalah: Komponen RevenueChart dan InsightChart menulis SVG <title> "
        "multi-line dengan beberapa expression. JSX mengubahnya jadi array "
        "6 elemen yang tidak bisa di-render React. Setiap render dashboard "
        "Analytics memunculkan error 'children prop of <title> tags ... "
        "found an Array with length 6'. Berdampak ke badge merah '1 Issue' "
        "di Next.js dev overlay dan error log berulang."
    )
    doc.add_paragraph(
        "Solusi: Konversi multi-expression ke template string tunggal. "
        "Sebelum: <title>{a} — {b} ({c} order)</title> menjadi "
        "<title>{`${a} — ${b} (${c} order)`}</title>. "
        "Sekalian fix warning Next.js scroll-behavior dengan menambah "
        "atribut data-scroll-behavior=\"smooth\" ke root <html>."
    )
    doc.add_paragraph(
        "File: src/components/RevenueChart.tsx, "
        "src/components/InsightChart.tsx, "
        "src/app/layout.tsx"
    )

    doc.add_page_break()

    # -- Section 8: Test Coverage --
    doc.add_heading("8. Test Coverage", level=1)

    doc.add_paragraph(
        "Total 98 unit test + 39 E2E/screenshot test. "
        "Semua CI workflow (lint, type-check, test, build) passing."
    )

    test_table = doc.add_table(rows=14, cols=3)
    test_table.style = "Light Shading Accent 1"
    for i, h in enumerate(["File", "Tests", "Cakupan"]):
        test_table.rows[0].cells[i].text = h

    test_data = [
        ("utils.test.ts", "7", "formatRupiah, generateOrderNumber, sanitizeText"),
        ("payment-mock.test.ts", "4", "QRIS payment, signature, webhook"),
        ("rate-limit.test.ts", "2", "Rate limiter allow/block"),
        ("notifications.test.ts", "3", "Phone normalization"),
        ("features.test.ts", "4", "Plan features matrix"),
        ("loyalty.test.ts", "9", "Points calculation + safe guards"),
        ("loyalty-tiers.test.ts", "17", "Tier ladder, multiplier, next-tier"),
        ("logo.test.ts", "6", "Logo config parser + defaults"),
        ("api-handler.test.ts", "8", "Error factories (400-429)"),
        ("theme-runtime.test.ts", "10", "Theme resolver, contrast, hex"),
        ("template-runtime.test.ts", "8", "Template resolver + button radius"),
        ("receipt.test.ts", "11", "Receipt formatter + WhatsApp wrap"),
        ("tax.test.ts", "9", "PPN calculator, EXCL/INCL, edge cases"),
    ]
    for r, (file, count, scope) in enumerate(test_data, start=1):
        test_table.rows[r].cells[0].text = file
        test_table.rows[r].cells[1].text = count
        test_table.rows[r].cells[2].text = scope

    doc.add_paragraph()

    # -- Footer --
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("— Akhir Laporan —")
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Save
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"✅ Report saved: {OUTPUT_PATH}")
    print(f"   Size: {os.path.getsize(OUTPUT_PATH) / 1024:.0f} KB")


if __name__ == "__main__":
    create_report()
